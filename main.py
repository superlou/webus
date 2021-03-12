#!/usr/bin/python3
import pandas as pd
from docx import Document
import subprocess


def generate_doc(df, filename):
    d = Document()

    for i, item in df.iterrows():
        if item.style == 'title':
            d.add_heading(item.text, 0)
        elif item.style == 'heading':
            parts = item.section.split('.')
            d.add_heading(item.text, len(parts) + 1)
        elif item.style == 'para':
            d.add_paragraph(item.text)

    d.save(filename)


def load_spreadsheet(filename):
    df = pd.read_excel(filename, dtype={
        'section': str,
        'style': str,
        'text': str,
        'trace': str,
    })

    df = df.fillna('')
    return df


def main():
    df = load_spreadsheet('reqs.xlsx')
    generate_doc(df, 'output.docx')
    # subprocess.run(['libreoffice', 'output.docx'], check=True)



if __name__ == '__main__':
    main()
